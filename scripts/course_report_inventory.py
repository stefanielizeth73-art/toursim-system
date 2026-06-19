import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")


def compact(text: str, limit: int = 180) -> str:
    text = " ".join((text or "").split())
    return text[:limit]


def inspect_docx(path: Path) -> None:
    try:
        doc = Document(str(path))
    except Exception as exc:
        print(f"### {path.relative_to(ROOT)}")
        print(f"ERROR: {exc}")
        return

    print(f"### {path.relative_to(ROOT)}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
    texts = [compact(p.text) for p in doc.paragraphs if compact(p.text)]
    for index, text in enumerate(texts[:60], 1):
        print(f"P{index}: {text}")
    for table_index, table in enumerate(doc.tables[:5]):
        print(f"TABLE {table_index}: rows={len(table.rows)} cols={len(table.columns)}")
        for row_index, row in enumerate(table.rows[:10]):
            cells = [compact(cell.text, 90) for cell in row.cells]
            print(f"  R{row_index}: {cells}")
    print()


def inspect_pdf(path: Path) -> None:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        print(f"### {path.relative_to(ROOT)}")
        print(f"ERROR: cannot import pypdf: {exc}")
        return

    reader = PdfReader(str(path))
    print(f"### {path.relative_to(ROOT)}")
    print(f"pages={len(reader.pages)}")
    for page_index, page in enumerate(reader.pages[:12], 1):
        text = page.extract_text() or ""
        text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
        print(f"--- PAGE {page_index} ---")
        print(text[:2500])
    print()


def main() -> None:
    for docx in sorted(DOCS.rglob("*.docx")):
        if docx.name.startswith("~$"):
            continue
        inspect_docx(docx)

    pdf = DOCS / "数据结构课程设计-2026.pdf"
    if pdf.exists():
        inspect_pdf(pdf)


if __name__ == "__main__":
    main()
